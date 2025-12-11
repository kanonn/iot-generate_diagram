# -*- coding: utf-8 -*-
"""
CloudFormation YAML から Word ドキュメントとアーキテクチャ図を同時生成
Word ドキュメントにアーキテクチャ図を挿入
"""

import os
import yaml
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from diagrams import Diagram, Cluster, Edge
from diagrams.aws.network import VPC, InternetGateway, PrivateSubnet, PublicSubnet, NATGateway, ELB, ALB, NLB, Route53, CF, APIGateway, VPCRouter
from diagrams.aws.compute import EC2, ECS, EKS, Lambda, Batch, ElasticBeanstalk
from diagrams.aws.database import RDS, Dynamodb, ElastiCache, Redshift, Neptune, Database
from diagrams.aws.storage import S3, EBS, EFS, FSx, Storage, Backup
from diagrams.aws.integration import SQS, SNS, Eventbridge, StepFunctions, MQ
from diagrams.aws.security import IAM, SecretsManager, KMS, WAF, Shield, CertificateManager
from diagrams.aws.management import Cloudwatch, SystemsManager, Cloudformation, Config
from diagrams.generic.blank import Blank


# ==================== CloudFormation YAML タグ処理 ====================

class CloudFormationYAMLLoader(yaml.SafeLoader):
    pass


def ref_constructor(loader, node):
    return {'Ref': loader.construct_scalar(node)}

def getatt_constructor(loader, node):
    if isinstance(node, yaml.ScalarNode):
        value = loader.construct_scalar(node)
        return {'Fn::GetAtt': value.split('.')}
    else:
        return {'Fn::GetAtt': loader.construct_sequence(node)}

def sub_constructor(loader, node):
    if isinstance(node, yaml.ScalarNode):
        return {'Fn::Sub': loader.construct_scalar(node)}
    else:
        return {'Fn::Sub': loader.construct_sequence(node)}

def join_constructor(loader, node):
    return {'Fn::Join': loader.construct_sequence(node)}

def select_constructor(loader, node):
    return {'Fn::Select': loader.construct_sequence(node)}

def getazs_constructor(loader, node):
    return {'Fn::GetAZs': loader.construct_scalar(node)}

def importvalue_constructor(loader, node):
    return {'Fn::ImportValue': loader.construct_scalar(node)}

def split_constructor(loader, node):
    return {'Fn::Split': loader.construct_sequence(node)}

def findinmap_constructor(loader, node):
    return {'Fn::FindInMap': loader.construct_sequence(node)}

def cidr_constructor(loader, node):
    return {'Fn::Cidr': loader.construct_sequence(node)}

def base64_constructor(loader, node):
    return {'Fn::Base64': loader.construct_scalar(node)}

def if_constructor(loader, node):
    return {'Fn::If': loader.construct_sequence(node)}

def equals_constructor(loader, node):
    return {'Fn::Equals': loader.construct_sequence(node)}

def not_constructor(loader, node):
    return {'Fn::Not': loader.construct_sequence(node)}

def and_constructor(loader, node):
    return {'Fn::And': loader.construct_sequence(node)}

def or_constructor(loader, node):
    return {'Fn::Or': loader.construct_sequence(node)}


CloudFormationYAMLLoader.add_constructor('!Ref', ref_constructor)
CloudFormationYAMLLoader.add_constructor('!GetAtt', getatt_constructor)
CloudFormationYAMLLoader.add_constructor('!Sub', sub_constructor)
CloudFormationYAMLLoader.add_constructor('!Join', join_constructor)
CloudFormationYAMLLoader.add_constructor('!Select', select_constructor)
CloudFormationYAMLLoader.add_constructor('!GetAZs', getazs_constructor)
CloudFormationYAMLLoader.add_constructor('!ImportValue', importvalue_constructor)
CloudFormationYAMLLoader.add_constructor('!Split', split_constructor)
CloudFormationYAMLLoader.add_constructor('!FindInMap', findinmap_constructor)
CloudFormationYAMLLoader.add_constructor('!Cidr', cidr_constructor)
CloudFormationYAMLLoader.add_constructor('!Base64', base64_constructor)
CloudFormationYAMLLoader.add_constructor('!If', if_constructor)
CloudFormationYAMLLoader.add_constructor('!Equals', equals_constructor)
CloudFormationYAMLLoader.add_constructor('!Not', not_constructor)
CloudFormationYAMLLoader.add_constructor('!And', and_constructor)
CloudFormationYAMLLoader.add_constructor('!Or', or_constructor)


def parse_yaml(yaml_file):
    try:
        with open(yaml_file, 'r', encoding='utf-8') as f:
            return yaml.load(f, Loader=CloudFormationYAMLLoader)
    except Exception as e:
        print(f"    Error: Failed to parse {yaml_file} - {e}")
        return None


# ==================== アイコンマッピング ====================

def get_icon_class(resource_type):
    icon_map = {
        'AWS::EC2::VPC': VPC,
        'AWS::EC2::Subnet': PrivateSubnet,
        'AWS::EC2::InternetGateway': InternetGateway,
        'AWS::EC2::VPCGatewayAttachment': InternetGateway,
        'AWS::EC2::NatGateway': NATGateway,
        'AWS::EC2::RouteTable': VPCRouter,
        'AWS::EC2::Route': VPCRouter,
        'AWS::EC2::SubnetRouteTableAssociation': VPCRouter,
        'AWS::EC2::SecurityGroup': VPCRouter,
        'AWS::ElasticLoadBalancingV2::LoadBalancer': ALB,
        'AWS::ElasticLoadBalancingV2::TargetGroup': ALB,
        'AWS::Route53::HostedZone': Route53,
        'AWS::EC2::Instance': EC2,
        'AWS::AutoScaling::AutoScalingGroup': EC2,
        'AWS::ECS::Cluster': ECS,
        'AWS::ECS::Service': ECS,
        'AWS::ECS::TaskDefinition': ECS,
        'AWS::EKS::Cluster': EKS,
        'AWS::Lambda::Function': Lambda,
        'AWS::Lambda::Permission': Lambda,
        'AWS::RDS::DBInstance': RDS,
        'AWS::RDS::DBCluster': RDS,
        'AWS::RDS::DBSubnetGroup': RDS,
        'AWS::DynamoDB::Table': Dynamodb,
        'AWS::S3::Bucket': S3,
        'AWS::EFS::FileSystem': EFS,
        'AWS::EFS::MountTarget': EFS,
        'AWS::EFS::AccessPoint': EFS,
        'AWS::Backup::BackupVault': Backup,
        'AWS::Backup::BackupPlan': Backup,
        'AWS::Backup::BackupSelection': Backup,
        'AWS::SQS::Queue': SQS,
        'AWS::SNS::Topic': SNS,
        'AWS::IAM::Role': IAM,
        'AWS::IAM::Policy': IAM,
        'AWS::CloudWatch::Alarm': Cloudwatch,
        'AWS::Logs::LogGroup': Cloudwatch,
        'AWS::Logs::MetricFilter': Cloudwatch,
        'AWS::SSM::Parameter': SystemsManager,
    }
    return icon_map.get(resource_type)


def extract_string_value(value):
    if isinstance(value, str):
        return value
    elif isinstance(value, dict):
        if 'Ref' in value:
            return f"!Ref {value['Ref']}"
        elif 'Fn::Sub' in value:
            sub_value = value['Fn::Sub']
            if isinstance(sub_value, str):
                return f"!Sub {sub_value}"
            return "!Sub [...]"
        elif 'Fn::GetAtt' in value:
            attrs = value['Fn::GetAtt']
            if isinstance(attrs, list):
                return f"!GetAtt {attrs[0]}.{attrs[1]}"
            return f"!GetAtt {attrs}"
        return str(value)
    return str(value)


def format_value_compact(value, max_length=100):
    if isinstance(value, dict):
        if 'Ref' in value:
            return f"!Ref {value['Ref']}"
        elif 'Fn::GetAtt' in value:
            attrs = value['Fn::GetAtt']
            if isinstance(attrs, list):
                return f"!GetAtt {attrs[0]}.{attrs[1]}"
            return f"!GetAtt {attrs}"
        elif 'Fn::Sub' in value:
            sub_value = value['Fn::Sub']
            if isinstance(sub_value, str):
                if len(sub_value) > max_length:
                    return f"!Sub {sub_value[:max_length-10]}..."
                return f"!Sub {sub_value}"
            return f"!Sub [...]"
        else:
            if not value:
                return "{}"
            items = []
            for k, v in list(value.items())[:3]:
                items.append(f"{k}: {format_value_compact(v, 30)}")
            result = "{" + ", ".join(items)
            if len(value) > 3:
                result += ", ..."
            result += "}"
            return result
    elif isinstance(value, list):
        if not value:
            return "[]"
        if len(value) == 1:
            return f"[{format_value_compact(value[0], 30)}]"
        if all(isinstance(v, (str, int, bool)) for v in value):
            if len(value) <= 3:
                return f"[{', '.join(str(v) for v in value)}]"
            else:
                return f"[{', '.join(str(v) for v in value[:3])}, ... ({len(value)} items)]"
        return f"[{len(value)} items]"
    elif isinstance(value, str):
        if len(value) > max_length:
            return f'"{value[:max_length-3]}..."'
        return f'"{value}"'
    else:
        return str(value)


def parse_xml(xml_string):
    from docx.oxml import parse_xml as docx_parse_xml
    return docx_parse_xml(xml_string)


def nsdecls(*prefixes):
    from docx.oxml.ns import nsdecls as docx_nsdecls
    return docx_nsdecls(*prefixes)


def add_heading_with_style(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in heading.runs:
        run.font.name = 'Arial'
        run.font.color.rgb = RGBColor(0, 51, 102)
    return heading


def get_resource_name(resource_info, resource_id):
    props = resource_info.get('Properties', {})
    tags = props.get('Tags', [])
    if tags:
        for tag in tags:
            if isinstance(tag, dict) and tag.get('Key') == 'Name':
                name_value = tag.get('Value')
                name = extract_string_value(name_value)
                if name and name != 'Name':
                    return name
    for key in ['FunctionName', 'DBInstanceIdentifier', 'BucketName', 
                'TableName', 'ClusterName', 'QueueName', 'TopicName', 'Name',
                'BackupVaultName', 'BackupPlanName', 'LogGroupName']:
        if key in props:
            value = props[key]
            name = extract_string_value(value)
            if name:
                return name
    return resource_id


def get_resource_label(resource_id, resource_data):
    props = resource_data.get('Properties', {})
    tags = props.get('Tags', [])
    if tags:
        for tag in tags:
            if isinstance(tag, dict) and tag.get('Key') == 'Name':
                name = extract_string_value(tag.get('Value'))
                if name and not name.startswith('Ref:'):
                    return name[:20]
    for key in ['FunctionName', 'DBInstanceIdentifier', 'BucketName', 
                'TableName', 'ClusterName', 'QueueName', 'TopicName', 'Name',
                'BackupVaultName', 'BackupPlanName', 'LogGroupName']:
        if key in props:
            name = extract_string_value(props[key])
            if name and not name.startswith('Ref:'):
                return name[:20]
    return resource_id[:15]


def flatten_dict(d, parent_key='', sep='.', max_depth=5, current_depth=0):
    items = []
    if current_depth >= max_depth:
        items.append((parent_key, str(d)))
        return items
    if isinstance(d, dict):
        for k, v in d.items():
            new_key = f"{parent_key}{sep}{k}" if parent_key else k
            if isinstance(v, dict):
                if any(key.startswith('Fn::') or key == 'Ref' for key in v.keys()):
                    items.append((new_key, format_value_compact(v)))
                else:
                    items.extend(flatten_dict(v, new_key, sep, max_depth, current_depth + 1))
            elif isinstance(v, list):
                if len(v) == 0:
                    items.append((new_key, '[]'))
                elif all(isinstance(item, (str, int, bool, type(None))) for item in v):
                    items.append((new_key, format_value_compact(v)))
                else:
                    for idx, item in enumerate(v):
                        if isinstance(item, dict):
                            items.extend(flatten_dict(item, f"{new_key}[{idx}]", sep, max_depth, current_depth + 1))
                        else:
                            items.append((f"{new_key}[{idx}]", format_value_compact(item)))
            else:
                items.append((new_key, format_value_compact(v)))
    else:
        items.append((parent_key, format_value_compact(d)))
    return items


def find_all_references(resources):
    relationships = []
    for source_id, source_data in resources.items():
        props = source_data.get('Properties', {})
        def find_refs(obj):
            refs = []
            if isinstance(obj, dict):
                if 'Ref' in obj:
                    refs.append(obj['Ref'])
                else:
                    for value in obj.values():
                        refs.extend(find_refs(value))
            elif isinstance(obj, list):
                for item in obj:
                    refs.extend(find_refs(item))
            return refs
        refs = find_refs(props)
        for target_id in refs:
            if target_id in resources:
                relationships.append({'from': source_id, 'to': target_id, 'type': 'ref'})
    return relationships


def categorize_resources(resources):
    categories = {
        'Network': [], 'Compute': [], 'Database': [], 'Storage': [],
        'Integration': [], 'Security': [], 'Management': [], 'Other': []
    }
    category_map = {
        'AWS::EC2::VPC': 'Network', 'AWS::EC2::Subnet': 'Network',
        'AWS::EC2::InternetGateway': 'Network', 'AWS::EC2::SecurityGroup': 'Security',
        'AWS::Lambda::Function': 'Compute', 'AWS::RDS::DBInstance': 'Database',
        'AWS::S3::Bucket': 'Storage', 'AWS::EFS::FileSystem': 'Storage',
        'AWS::EFS::MountTarget': 'Storage', 'AWS::EFS::AccessPoint': 'Storage',
        'AWS::Backup::BackupVault': 'Storage', 'AWS::Backup::BackupPlan': 'Storage',
        'AWS::IAM::Role': 'Security', 'AWS::Logs::LogGroup': 'Management',
        'AWS::Logs::MetricFilter': 'Management',
    }
    for resource_id, resource_data in resources.items():
        resource_type = resource_data.get('Type', '')
        category = category_map.get(resource_type, 'Other')
        categories[category].append((resource_id, resource_data, resource_type))
    return {k: v for k, v in categories.items() if v}


def generate_diagram(yaml_file, output_dir):
    """アーキテクチャ図を生成"""
    template = parse_yaml(yaml_file)
    if not template or 'Resources' not in template:
        return None
    
    resources = template['Resources']
    base_name = os.path.splitext(os.path.basename(yaml_file))[0]
    output_filename = os.path.join(output_dir, f'{base_name}_diagram')
    
    relationships = find_all_references(resources)
    categories = categorize_resources(resources)
    
    graph_attr = {"fontsize": "11", "bgcolor": "white", "pad": "0.5", "nodesep": "0.8", "ranksep": "1.0"}
    
    try:
        with Diagram(
            f"{base_name}",
            filename=output_filename,
            show=False,
            direction="TB",
            outformat="png",
            graph_attr=graph_attr
        ):
            nodes = {}
            for category, resource_list in categories.items():
                with Cluster(f"{category} ({len(resource_list)})"):
                    for resource_id, resource_data, resource_type in resource_list:
                        icon_class = get_icon_class(resource_type)
                        label = get_resource_label(resource_id, resource_data)
                        if icon_class:
                            node = icon_class(label)
                        else:
                            node = Blank(label)
                        nodes[resource_id] = node
            
            for rel in relationships:
                if rel['from'] in nodes and rel['to'] in nodes:
                    nodes[rel['from']] >> Edge(color="blue", style="solid") >> nodes[rel['to']]
        
        return f"{output_filename}.png"
    except Exception as e:
        print(f"    Warning: Failed to generate diagram - {e}")
        return None


def generate_word_with_diagram(yaml_file, output_dir='docs'):
    """Word ドキュメントとアーキテクチャ図を同時生成"""
    
    template = parse_yaml(yaml_file)
    if not template or 'Resources' not in template:
        print(f"  Skip: {yaml_file} - No resources found")
        return None
    
    os.makedirs(output_dir, exist_ok=True)
    
    # アーキテクチャ図を先に生成
    print(f"  Generating diagram...")
    diagram_path = generate_diagram(yaml_file, output_dir)
    
    all_resources = template['Resources']
    template_description = template.get('Description', '')
    template_version = template.get('AWSTemplateFormatVersion', '')
    parameters = template.get('Parameters', {})
    mappings = template.get('Mappings', {})
    conditions = template.get('Conditions', {})
    outputs = template.get('Outputs', {})
    
    yaml_basename = os.path.splitext(os.path.basename(yaml_file))[0]
    
    # Word ドキュメント作成
    doc = Document()
    core_properties = doc.core_properties
    core_properties.author = 'CloudFormation Documentation Generator'
    core_properties.title = yaml_basename
    core_properties.created = datetime.now()
    
    # Title
    title = doc.add_heading(yaml_basename, level=1)
    for run in title.runs:
        run.font.name = 'Arial'
        run.font.color.rgb = RGBColor(0, 51, 102)
    
    # Template Information
    add_heading_with_style(doc, 'Template Information', level=2)
    info_table = doc.add_table(rows=2, cols=2)
    info_table.style = 'Light Grid Accent 1'
    info_table.rows[0].cells[0].text = 'Property'
    info_table.rows[0].cells[1].text = 'Value'
    for cell in info_table.rows[0].cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), '4472C4'))
        cell._element.get_or_add_tcPr().append(shading_elm)
    info_table.rows[1].cells[0].text = 'AWSTemplateFormatVersion'
    info_table.rows[1].cells[1].text = template_version if template_version else 'N/A'
    doc.add_paragraph()
    
    # Description
    if template_description:
        add_heading_with_style(doc, 'Description', level=2)
        doc.add_paragraph(template_description)
        doc.add_paragraph()
    
    # ==================== Architecture Diagram ====================
    if diagram_path and os.path.exists(diagram_path):
        add_heading_with_style(doc, 'Architecture Diagram', level=2)
        doc.add_paragraph('The following diagram illustrates the architecture of this CloudFormation template:')
        doc.add_paragraph()
        
        try:
            # 画像を挿入（幅 6 インチ）
            doc.add_picture(diagram_path, width=Inches(6))
            doc.add_paragraph()
        except Exception as e:
            print(f"    Warning: Failed to insert diagram - {e}")
            doc.add_paragraph(f'[Diagram file: {os.path.basename(diagram_path)}]')
            doc.add_paragraph()
    
    # Parameters
    if parameters:
        add_heading_with_style(doc, 'Parameters', level=2)
        doc.add_paragraph(f'This template defines {len(parameters)} parameter(s).')
        doc.add_paragraph()
        for param_name, param_data in parameters.items():
            add_heading_with_style(doc, param_name, level=3)
            param_items = flatten_dict(param_data)
            if param_items:
                param_table = doc.add_table(rows=len(param_items) + 1, cols=2)
                param_table.style = 'Light Grid Accent 1'
                param_table.rows[0].cells[0].text = 'Property'
                param_table.rows[0].cells[1].text = 'Value'
                for cell in param_table.rows[0].cells:
                    cell.paragraphs[0].runs[0].font.bold = True
                    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), '4472C4'))
                    cell._element.get_or_add_tcPr().append(shading_elm)
                for idx, (key, value) in enumerate(param_items, start=1):
                    param_table.rows[idx].cells[0].text = key
                    param_table.rows[idx].cells[1].text = str(value)
            doc.add_paragraph()
    
    # Mappings
    if mappings:
        add_heading_with_style(doc, 'Mappings', level=2)
        doc.add_paragraph(f'This template defines {len(mappings)} mapping(s).')
        doc.add_paragraph()
        for mapping_name, mapping_data in mappings.items():
            add_heading_with_style(doc, mapping_name, level=3)
            mapping_items = flatten_dict(mapping_data)
            if mapping_items:
                mapping_table = doc.add_table(rows=len(mapping_items) + 1, cols=2)
                mapping_table.style = 'Light Grid Accent 1'
                mapping_table.rows[0].cells[0].text = 'Key Path'
                mapping_table.rows[0].cells[1].text = 'Value'
                for cell in mapping_table.rows[0].cells:
                    cell.paragraphs[0].runs[0].font.bold = True
                    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), '4472C4'))
                    cell._element.get_or_add_tcPr().append(shading_elm)
                for idx, (key, value) in enumerate(mapping_items, start=1):
                    mapping_table.rows[idx].cells[0].text = key
                    mapping_table.rows[idx].cells[1].text = str(value)
            doc.add_paragraph()
    
    # Conditions
    if conditions:
        add_heading_with_style(doc, 'Conditions', level=2)
        doc.add_paragraph(f'This template defines {len(conditions)} condition(s).')
        doc.add_paragraph()
        cond_table = doc.add_table(rows=len(conditions) + 1, cols=2)
        cond_table.style = 'Light Grid Accent 1'
        cond_table.rows[0].cells[0].text = 'Condition Name'
        cond_table.rows[0].cells[1].text = 'Expression'
        for cell in cond_table.rows[0].cells:
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), '4472C4'))
            cell._element.get_or_add_tcPr().append(shading_elm)
        for idx, (cond_name, cond_expr) in enumerate(conditions.items(), start=1):
            cond_table.rows[idx].cells[0].text = cond_name
            cond_table.rows[idx].cells[1].text = format_value_compact(cond_expr, 200)
        doc.add_paragraph()
    
    # Resources
    if all_resources:
        add_heading_with_style(doc, 'Resources', level=2)
        doc.add_paragraph(f'This template contains {len(all_resources)} resource(s).')
        doc.add_paragraph()
        
        overview_table = doc.add_table(rows=len(all_resources) + 1, cols=3)
        overview_table.style = 'Light Grid Accent 1'
        overview_table.rows[0].cells[0].text = 'Resource ID'
        overview_table.rows[0].cells[1].text = 'Resource Name'
        overview_table.rows[0].cells[2].text = 'Type'
        for cell in overview_table.rows[0].cells:
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), '4472C4'))
            cell._element.get_or_add_tcPr().append(shading_elm)
        
        for idx, (resource_id, resource_data) in enumerate(all_resources.items(), start=1):
            resource_type = resource_data.get('Type', 'Unknown')
            resource_name = get_resource_name(resource_data, resource_id)
            overview_table.rows[idx].cells[0].text = resource_id
            overview_table.rows[idx].cells[1].text = resource_name
            overview_table.rows[idx].cells[2].text = resource_type
        
        doc.add_paragraph()
        doc.add_page_break()
        
        # Resource Details
        add_heading_with_style(doc, 'Resource Details', level=2)
        doc.add_paragraph()
        
        for resource_idx, (resource_id, resource_data) in enumerate(all_resources.items(), start=1):
            resource_type = resource_data.get('Type', 'Unknown')
            resource_props = resource_data.get('Properties', {})
            resource_name = get_resource_name(resource_data, resource_id)
            
            add_heading_with_style(doc, f'{resource_idx}. {resource_name}', level=3)
            
            info_table = doc.add_table(rows=3, cols=2)
            info_table.style = 'Light Grid Accent 1'
            info_table.rows[0].cells[0].text = 'Property'
            info_table.rows[0].cells[1].text = 'Value'
            for cell in info_table.rows[0].cells:
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), '4472C4'))
                cell._element.get_or_add_tcPr().append(shading_elm)
            
            info_table.rows[1].cells[0].text = 'Resource ID'
            info_table.rows[1].cells[1].text = resource_id
            info_table.rows[2].cells[0].text = 'Resource Type'
            info_table.rows[2].cells[1].text = resource_type
            doc.add_paragraph()
            
            if resource_props:
                add_heading_with_style(doc, 'Properties', level=4)
                flattened = flatten_dict(resource_props)
                if flattened:
                    prop_table = doc.add_table(rows=len(flattened) + 1, cols=2)
                    prop_table.style = 'Light Grid Accent 1'
                    prop_table.rows[0].cells[0].text = 'Property Path'
                    prop_table.rows[0].cells[1].text = 'Value'
                    for cell in prop_table.rows[0].cells:
                        cell.paragraphs[0].runs[0].font.bold = True
                        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                        shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), '4472C4'))
                        cell._element.get_or_add_tcPr().append(shading_elm)
                    for idx, (key, value) in enumerate(flattened, start=1):
                        prop_table.rows[idx].cells[0].text = key
                        prop_table.rows[idx].cells[1].text = str(value)
            
            doc.add_paragraph()
            
            if resource_idx < len(all_resources):
                doc.add_paragraph('_' * 80)
                doc.add_paragraph()
    
    # Outputs
    if outputs:
        doc.add_page_break()
        add_heading_with_style(doc, 'Outputs', level=2)
        doc.add_paragraph(f'This template defines {len(outputs)} output(s).')
        doc.add_paragraph()
        for output_name, output_data in outputs.items():
            add_heading_with_style(doc, output_name, level=3)
            output_items = flatten_dict(output_data)
            if output_items:
                output_table = doc.add_table(rows=len(output_items) + 1, cols=2)
                output_table.style = 'Light Grid Accent 1'
                output_table.rows[0].cells[0].text = 'Property'
                output_table.rows[0].cells[1].text = 'Value'
                for cell in output_table.rows[0].cells:
                    cell.paragraphs[0].runs[0].font.bold = True
                    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), '4472C4'))
                    cell._element.get_or_add_tcPr().append(shading_elm)
                for idx, (key, value) in enumerate(output_items, start=1):
                    output_table.rows[idx].cells[0].text = key
                    output_table.rows[idx].cells[1].text = str(value)
            doc.add_paragraph()
    
    # Save
    safe_name = yaml_basename.replace('/', '-').replace('\\', '-').replace(':', '-').replace('*', '-').replace('?', '-').replace('"', '-').replace('<', '-').replace('>', '-').replace('|', '-')
    if len(safe_name) > 50:
        safe_name = safe_name[:47] + "..."
    
    output_file = os.path.join(output_dir, f'{safe_name}.docx')
    doc.save(output_file)
    
    return output_file


def generate_all_docs_with_diagrams(input_dir='aws-resources', output_dir='docs-with-diagrams'):
    """すべての YAML ファイルからドキュメントと図を生成"""
    
    print("="*80)
    print("CloudFormation Documentation Generator (With Architecture Diagrams)")
    print("="*80)
    print(f"\nInput directory: {input_dir}")
    print(f"Output directory: {output_dir}\n")
    
    yaml_files = []
    for root, dirs, files in os.walk(input_dir):
        for file in files:
            if file.endswith('.yaml') or file.endswith('.yml'):
                yaml_files.append(os.path.join(root, file))
    
    print(f"Found {len(yaml_files)} YAML file(s)\n")
    
    success_count = 0
    error_count = 0
    
    for yaml_file in yaml_files:
        print(f"Processing: {os.path.basename(yaml_file)}")
        
        try:
            output_file = generate_word_with_diagram(yaml_file, output_dir)
            if output_file:
                print(f"  -> Generated: {os.path.basename(output_file)}")
                success_count += 1
            else:
                error_count += 1
        except Exception as e:
            print(f"  -> Error: {e}")
            import traceback
            traceback.print_exc()
            error_count += 1
    
    print("\n" + "="*80)
    print(f"Complete!")
    print(f"  Success: {success_count} document(s)")
    print(f"  Errors: {error_count} file(s)")
    print(f"Output directory: {os.path.abspath(output_dir)}")
    print("="*80)


def main():
    import argparse
    
    parser = argparse.ArgumentParser(description='Generate Word documentation with embedded architecture diagrams')
    parser.add_argument('--input-dir', default='aws-resources', help='Input directory containing YAML files')
    parser.add_argument('--output-dir', default='docs-with-diagrams', help='Output directory')
    
    args = parser.parse_args()
    
    generate_all_docs_with_diagrams(args.input_dir, args.output_dir)


if __name__ == '__main__':
    main()