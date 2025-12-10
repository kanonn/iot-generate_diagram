# -*- coding: utf-8 -*-
"""
修复版 - 完整支持 CloudFormation YAML 标签
"""

import os
import yaml
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


# ==================== CloudFormation YAML 标签处理 ====================

class CloudFormationYAMLLoader(yaml.SafeLoader):
    """自定义 YAML Loader 支持 CloudFormation 标签"""
    pass


def ref_constructor(loader, node):
    return {'Ref': loader.construct_scalar(node)}

def getatt_constructor(loader, node):
    if isinstance(node, yaml.ScalarNode):
        value = loader.construct_scalar(node)
        return {'Fn::GetAtt': value.split('.')}
    else:
        return {'Fn::GetAtt': loader.construct_sequence(node)}

def sub_constructor(loader, node):
    if isinstance(node, yaml.ScalarNode):
        return {'Fn::Sub': loader.construct_scalar(node)}
    else:
        return {'Fn::Sub': loader.construct_sequence(node)}

def join_constructor(loader, node):
    return {'Fn::Join': loader.construct_sequence(node)}

def select_constructor(loader, node):
    return {'Fn::Select': loader.construct_sequence(node)}

def getazs_constructor(loader, node):
    return {'Fn::GetAZs': loader.construct_scalar(node)}

def importvalue_constructor(loader, node):
    return {'Fn::ImportValue': loader.construct_scalar(node)}

def split_constructor(loader, node):
    return {'Fn::Split': loader.construct_sequence(node)}

def findinmap_constructor(loader, node):
    return {'Fn::FindInMap': loader.construct_sequence(node)}

def cidr_constructor(loader, node):
    return {'Fn::Cidr': loader.construct_sequence(node)}

def base64_constructor(loader, node):
    return {'Fn::Base64': loader.construct_scalar(node)}

def if_constructor(loader, node):
    return {'Fn::If': loader.construct_sequence(node)}

def equals_constructor(loader, node):
    return {'Fn::Equals': loader.construct_sequence(node)}

def not_constructor(loader, node):
    return {'Fn::Not': loader.construct_sequence(node)}

def and_constructor(loader, node):
    return {'Fn::And': loader.construct_sequence(node)}

def or_constructor(loader, node):
    return {'Fn::Or': loader.construct_sequence(node)}


# 注册所有标签
CloudFormationYAMLLoader.add_constructor('!Ref', ref_constructor)
CloudFormationYAMLLoader.add_constructor('!GetAtt', getatt_constructor)
CloudFormationYAMLLoader.add_constructor('!Sub', sub_constructor)
CloudFormationYAMLLoader.add_constructor('!Join', join_constructor)
CloudFormationYAMLLoader.add_constructor('!Select', select_constructor)
CloudFormationYAMLLoader.add_constructor('!GetAZs', getazs_constructor)
CloudFormationYAMLLoader.add_constructor('!ImportValue', importvalue_constructor)
CloudFormationYAMLLoader.add_constructor('!Split', split_constructor)
CloudFormationYAMLLoader.add_constructor('!FindInMap', findinmap_constructor)
CloudFormationYAMLLoader.add_constructor('!Cidr', cidr_constructor)
CloudFormationYAMLLoader.add_constructor('!Base64', base64_constructor)
CloudFormationYAMLLoader.add_constructor('!If', if_constructor)
CloudFormationYAMLLoader.add_constructor('!Equals', equals_constructor)
CloudFormationYAMLLoader.add_constructor('!Not', not_constructor)
CloudFormationYAMLLoader.add_constructor('!And', and_constructor)
CloudFormationYAMLLoader.add_constructor('!Or', or_constructor)


def parse_yaml(yaml_file):
    """解析 CloudFormation YAML 文件"""
    try:
        with open(yaml_file, 'r', encoding='utf-8') as f:
            return yaml.load(f, Loader=CloudFormationYAMLLoader)
    except Exception as e:
        print(f"    Error parsing {yaml_file}: {e}")
        return None


# ==================== 辅助函数 ====================

def extract_string_value(value):
    """从可能包含内置函数的值中提取字符串"""
    if isinstance(value, str):
        return value
    elif isinstance(value, dict):
        if 'Ref' in value:
            return f"Ref:{value['Ref']}"
        elif 'Fn::Sub' in value:
            sub_value = value['Fn::Sub']
            if isinstance(sub_value, str):
                return sub_value
            else:
                return str(sub_value[0]) if sub_value else "Sub:..."
        elif 'Fn::Join' in value:
            return "Join:..."
        elif 'Fn::GetAtt' in value:
            attrs = value['Fn::GetAtt']
            if isinstance(attrs, list):
                return f"GetAtt:{attrs[0]}.{attrs[1]}"
            else:
                return f"GetAtt:{attrs}"
        else:
            return str(value)
    else:
        return str(value)


def format_intrinsic_function(value, indent=0):
    """格式化 CloudFormation 内置函数显示"""
    indent_str = "  " * indent
    
    if isinstance(value, dict):
        if 'Ref' in value:
            return f"!Ref {value['Ref']}"
        elif 'Fn::GetAtt' in value:
            attrs = value['Fn::GetAtt']
            if isinstance(attrs, list):
                return f"!GetAtt {attrs[0]}.{attrs[1]}"
            else:
                return f"!GetAtt {attrs}"
        elif 'Fn::Sub' in value:
            sub_value = value['Fn::Sub']
            if isinstance(sub_value, str):
                if len(sub_value) > 60:
                    return f"!Sub '{sub_value[:57]}...'"
                return f"!Sub '{sub_value}'"
            else:
                return f"!Sub [{sub_value[0] if sub_value else '...'}]"
        elif 'Fn::Join' in value:
            join_parts = value['Fn::Join']
            delimiter = join_parts[0]
            return f"!Join ['{delimiter}', [...]]"
        elif 'Fn::Select' in value:
            select_parts = value['Fn::Select']
            return f"!Select [{select_parts[0]}, ...]"
        elif 'Fn::GetAZs' in value:
            return f"!GetAZs {value['Fn::GetAZs']}"
        elif 'Fn::ImportValue' in value:
            return f"!ImportValue {value['Fn::ImportValue']}"
        elif 'Fn::Split' in value:
            return f"!Split [{value['Fn::Split'][0]}, ...]"
        elif 'Fn::FindInMap' in value:
            return f"!FindInMap [{value['Fn::FindInMap'][0]}, ...]"
        elif 'Fn::Cidr' in value:
            return f"!Cidr [...]"
        elif 'Fn::Base64' in value:
            return f"!Base64 ..."
        elif 'Fn::If' in value:
            return f"!If [Condition, TrueValue, FalseValue]"
        else:
            if not value:
                return "{}"
            lines = ["{"]
            for k, v in value.items():
                formatted = format_intrinsic_function(v, indent + 1)
                lines.append(f"{indent_str}  {k}: {formatted}")
            lines.append(f"{indent_str}}}")
            return "\n".join(lines)
    
    elif isinstance(value, list):
        if not value:
            return "[]"
        if len(value) == 1:
            return f"[{format_intrinsic_function(value[0])}]"
        
        simple = all(isinstance(v, (str, int, bool)) for v in value)
        if simple and len(value) <= 3:
            return f"[{', '.join(str(v) for v in value)}]"
        
        lines = ["["]
        for item in value:
            formatted = format_intrinsic_function(item, indent + 1)
            lines.append(f"{indent_str}  - {formatted}")
        lines.append(f"{indent_str}]")
        return "\n".join(lines)
    
    elif isinstance(value, str):
        if len(value) > 60:
            return f'"{value[:57]}..."'
        return f'"{value}"'
    
    else:
        return str(value)


def parse_xml(xml_string):
    from docx.oxml import parse_xml as docx_parse_xml
    return docx_parse_xml(xml_string)


def nsdecls(*prefixes):
    from docx.oxml.ns import nsdecls as docx_nsdecls
    return docx_nsdecls(*prefixes)


def add_heading_with_style(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    for run in heading.runs:
        run.font.name = 'Microsoft YaHei'
        run.font.color.rgb = RGBColor(0, 51, 102)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    return heading


def add_code_block(doc, code_text):
    paragraph = doc.add_paragraph()
    paragraph.style = 'Normal'
    
    run = paragraph.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
    
    paragraph.paragraph_format.left_indent = Inches(0.3)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    
    return paragraph


def get_resource_name(resource_info, resource_id):
    """获取资源名称（修复版）"""
    props = resource_info.get('Properties', {})
    
    # 尝试从 Tags 获取
    tags = props.get('Tags', [])
    if tags:
        for tag in tags:
            if isinstance(tag, dict) and tag.get('Key') == 'Name':
                name_value = tag.get('Value')
                # 提取字符串值
                name = extract_string_value(name_value)
                if name and name != 'Name':
                    return name
    
    # 从其他属性获取
    for key in ['FunctionName', 'DBInstanceIdentifier', 'BucketName', 
                'TableName', 'ClusterName', 'QueueName', 'TopicName', 'Name']:
        if key in props:
            value = props[key]
            name = extract_string_value(value)
            if name:
                return name
    
    # 如果都没有，返回 resource_id
    return resource_id


def get_resource_description(resource_type):
    descriptions = {
        'AWS::EC2::VPC': 'Amazon Virtual Private Cloud - 虚拟私有云，提供隔离的网络环境',
        'AWS::EC2::Subnet': '子网 - VPC 内的网络分段，可以是公有或私有',
        'AWS::EC2::InternetGateway': '互联网网关 - 允许 VPC 与互联网通信',
        'AWS::EC2::VPCGatewayAttachment': 'VPC 网关连接 - 将 Internet Gateway 连接到 VPC',
        'AWS::EC2::RouteTable': '路由表 - 定义网络流量的路由规则',
        'AWS::EC2::Route': '路由 - 路由表中的单条路由规则',
        'AWS::EC2::SubnetRouteTableAssociation': '子网路由表关联 - 将路由表关联到子网',
        'AWS::EC2::NatGateway': 'NAT 网关 - 允许私有子网中的资源访问互联网',
        'AWS::EC2::SecurityGroup': '安全组 - 虚拟防火墙，控制入站和出站流量',
        'AWS::EC2::Instance': 'EC2 实例 - 虚拟服务器',
        'AWS::ECS::Cluster': 'ECS 集群 - 容器编排服务集群',
        'AWS::Lambda::Function': 'Lambda 函数 - 无服务器计算服务',
        'AWS::S3::Bucket': 'S3 存储桶 - 对象存储服务',
        'AWS::IAM::Role': 'IAM 角色 - 身份和访问管理角色',
    }
    
    return descriptions.get(resource_type, 'AWS 资源')


def generate_word_document(yaml_file, output_dir='docs'):
    """为单个 YAML 文件生成 Word 文档"""
    
    template = parse_yaml(yaml_file)
    if not template or 'Resources' not in template:
        print(f"  Skipping {yaml_file} - No resources found")
        return None
    
    os.makedirs(output_dir, exist_ok=True)
    
    resource_id = list(template['Resources'].keys())[0]
    resource_data = template['Resources'][resource_id]
    resource_type = resource_data.get('Type', 'Unknown')
    resource_props = resource_data.get('Properties', {})
    
    # 修复：获取资源名称（处理可能是 dict 的情况）
    resource_name = get_resource_name(resource_data, resource_id)
    
    doc = Document()
    
    core_properties = doc.core_properties
    core_properties.author = 'AWS Documentation Generator'
    core_properties.title = f'{resource_name} - Technical Documentation'
    core_properties.created = datetime.now()
    
    # 封面
    title = doc.add_heading(f'{resource_name}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(28)
        run.font.color.rgb = RGBColor(0, 51, 102)
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    subtitle = doc.add_paragraph('Technical Documentation')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_paragraph()
    
    type_para = doc.add_paragraph(f'Resource Type: {resource_type}')
    type_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in type_para.runs:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 102, 204)
    
    doc.add_paragraph()
    
    date_para = doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in date_para.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(150, 150, 150)
    
    doc.add_page_break()
    
    # 目录
    add_heading_with_style(doc, 'Table of Contents', level=1)
    
    toc_items = [
        '1. Overview',
        '2. Resource Details',
        '3. Properties',
        '4. Dependencies & References',
        '5. Tags',
        '6. YAML Source Code'
    ]
    
    for item in toc_items:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_page_break()
    
    # 1. Overview
    add_heading_with_style(doc, '1. Overview', level=1)
    
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Light Grid Accent 1'
    
    table.rows[0].cells[0].text = 'Property'
    table.rows[0].cells[1].text = 'Value'
    
    for cell in table.rows[0].cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), '4472C4'))
        cell._element.get_or_add_tcPr().append(shading_elm)
    
    table.rows[1].cells[0].text = 'Resource Name'
    table.rows[1].cells[1].text = resource_name
    
    table.rows[2].cells[0].text = 'Resource ID'
    table.rows[2].cells[1].text = resource_id
    
    table.rows[3].cells[0].text = 'Resource Type'
    table.rows[3].cells[1].text = resource_type
    
    doc.add_paragraph()
    
    add_heading_with_style(doc, 'Description', level=2)
    description = get_resource_description(resource_type)
    doc.add_paragraph(description)
    
    doc.add_paragraph()
    
    # 3. Properties
    add_heading_with_style(doc, '3. Properties', level=1)
    
    if resource_props:
        doc.add_paragraph('This resource has the following properties configured:')
        doc.add_paragraph()
        
        prop_table = doc.add_table(rows=len(resource_props) + 1, cols=3)
        prop_table.style = 'Light Grid Accent 1'
        
        prop_table.rows[0].cells[0].text = 'Property Name'
        prop_table.rows[0].cells[1].text = 'Type'
        prop_table.rows[0].cells[2].text = 'Value'
        
        for cell in prop_table.rows[0].cells:
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), '4472C4'))
            cell._element.get_or_add_tcPr().append(shading_elm)
        
        for idx, (prop_name, prop_value) in enumerate(resource_props.items(), start=1):
            prop_table.rows[idx].cells[0].text = prop_name
            
            # 类型判断
            if isinstance(prop_value, dict) and any(k.startswith('Fn::') or k == 'Ref' for k in prop_value.keys()):
                value_type = 'Intrinsic Function'
            elif isinstance(prop_value, dict):
                value_type = 'Object'
            elif isinstance(prop_value, list):
                value_type = 'Array'
            elif isinstance(prop_value, str):
                value_type = 'String'
            elif isinstance(prop_value, int):
                value_type = 'Integer'
            elif isinstance(prop_value, bool):
                value_type = 'Boolean'
            else:
                value_type = 'Unknown'
            
            prop_table.rows[idx].cells[1].text = value_type
            
            formatted_value = format_intrinsic_function(prop_value)
            prop_table.rows[idx].cells[2].text = formatted_value
    else:
        doc.add_paragraph('No properties configured for this resource.')
    
    doc.add_paragraph()
    
    # 4. Dependencies & References
    add_heading_with_style(doc, '4. Dependencies & References', level=1)
    
    refs = []
    getattrs = []
    
    def find_references(obj, path=""):
        if isinstance(obj, dict):
            if 'Ref' in obj:
                refs.append((path, obj['Ref']))
            elif 'Fn::GetAtt' in obj:
                getattrs.append((path, obj['Fn::GetAtt']))
            else:
                for key, value in obj.items():
                    find_references(value, f"{path}.{key}" if path else key)
        elif isinstance(obj, list):
            for idx, item in enumerate(obj):
                find_references(item, f"{path}[{idx}]")
    
    find_references(resource_props)
    
    if refs:
        add_heading_with_style(doc, 'References (!Ref)', level=2)
        ref_table = doc.add_table(rows=len(refs) + 1, cols=2)
        ref_table.style = 'Light Grid Accent 1'
        
        ref_table.rows[0].cells[0].text = 'Property Path'
        ref_table.rows[0].cells[1].text = 'Referenced Resource'
        
        for cell in ref_table.rows[0].cells:
            cell.paragraphs[0].runs[0].font.bold = True
        
        for idx, (path, ref) in enumerate(refs, start=1):
            ref_table.rows[idx].cells[0].text = path
            ref_table.rows[idx].cells[1].text = ref
        
        doc.add_paragraph()
    
    if getattrs:
        add_heading_with_style(doc, 'Attribute References (!GetAtt)', level=2)
        getatt_table = doc.add_table(rows=len(getattrs) + 1, cols=2)
        getatt_table.style = 'Light Grid Accent 1'
        
        getatt_table.rows[0].cells[0].text = 'Property Path'
        getatt_table.rows[0].cells[1].text = 'Attribute Reference'
        
        for cell in getatt_table.rows[0].cells:
            cell.paragraphs[0].runs[0].font.bold = True
        
        for idx, (path, getatt) in enumerate(getattrs, start=1):
            getatt_table.rows[idx].cells[0].text = path
            if isinstance(getatt, list):
                getatt_table.rows[idx].cells[1].text = f"{getatt[0]}.{getatt[1]}"
            else:
                getatt_table.rows[idx].cells[1].text = str(getatt)
        
        doc.add_paragraph()
    
    if not refs and not getattrs:
        doc.add_paragraph('This resource has no explicit references to other resources.')
    
    # 5. Tags
    add_heading_with_style(doc, '5. Tags', level=1)
    
    tags = resource_props.get('Tags', [])
    
    if tags:
        tag_table = doc.add_table(rows=len(tags) + 1, cols=2)
        tag_table.style = 'Light Grid Accent 1'
        
        tag_table.rows[0].cells[0].text = 'Key'
        tag_table.rows[0].cells[1].text = 'Value'
        
        for cell in tag_table.rows[0].cells:
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), '4472C4'))
            cell._element.get_or_add_tcPr().append(shading_elm)
        
        for idx, tag in enumerate(tags, start=1):
            if isinstance(tag, dict):
                key = extract_string_value(tag.get('Key', ''))
                value = extract_string_value(tag.get('Value', ''))
                tag_table.rows[idx].cells[0].text = key
                tag_table.rows[idx].cells[1].text = value
    else:
        doc.add_paragraph('No tags configured for this resource.')
    
    doc.add_paragraph()
    
    # 6. YAML Source Code
    add_heading_with_style(doc, '6. YAML Source Code', level=1)
    
    doc.add_paragraph('Complete YAML definition for this resource:')
    doc.add_paragraph()
    
    with open(yaml_file, 'r', encoding='utf-8') as f:
        yaml_content = f.read()
    
    add_code_block(doc, yaml_content)
    
    # 保存
    safe_name = resource_name.replace('/', '-').replace('\\', '-').replace(':', '-').replace('*', '-').replace('?', '-').replace('"', '-').replace('<', '-').replace('>', '-').replace('|', '-')
    # 限制文件名长度
    if len(safe_name) > 50:
        safe_name = safe_name[:47] + "..."
    
    output_file = os.path.join(output_dir, f'{safe_name}.docx')
    
    doc.save(output_file)
    
    return output_file


def generate_all_docs(input_dir='aws-resources', output_dir='aws-docs'):
    """为所有 YAML 文件生成文档"""
    
    print("="*80)
    print("AWS CloudFormation Documentation Generator (Fixed)")
    print("="*80)
    print(f"\nInput directory: {input_dir}")
    print(f"Output directory: {output_dir}\n")
    
    yaml_files = []
    for root, dirs, files in os.walk(input_dir):
        for file in files:
            if file.endswith('.yaml') or file.endswith('.yml'):
                yaml_files.append(os.path.join(root, file))
    
    print(f"Found {len(yaml_files)} YAML files\n")
    
    success_count = 0
    error_count = 0
    
    for yaml_file in yaml_files:
        print(f"Processing: {os.path.basename(yaml_file)}")
        
        try:
            output_file = generate_word_document(yaml_file, output_dir)
            if output_file:
                print(f"  -> Generated: {os.path.basename(output_file)}")
                success_count += 1
            else:
                error_count += 1
        except Exception as e:
            print(f"  -> Error: {e}")
            import traceback
            traceback.print_exc()
            error_count += 1
    
    print("\n" + "="*80)
    print(f"Complete!")
    print(f"  Success: {success_count} documents")
    print(f"  Errors: {error_count} files")
    print(f"Output directory: {os.path.abspath(output_dir)}")
    print("="*80)


def main():
    import argparse
    
    parser = argparse.ArgumentParser(description='Generate Word documentation from CloudFormation YAML files')
    parser.add_argument('--input-dir', default='aws-resources', help='Input directory containing YAML files')
    parser.add_argument('--output-dir', default='aws-docs', help='Output directory for Word documents')
    
    args = parser.parse_args()
    
    generate_all_docs(args.input_dir, args.output_dir)


if __name__ == '__main__':
    main()